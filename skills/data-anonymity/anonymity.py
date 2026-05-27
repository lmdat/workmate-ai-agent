#!/usr/bin/env python3
"""
Data Anonymity — xử lý local, không có network call.
Hỗ trợ: .txt, .csv, .docx, .xlsx, .pdf (convert → .docx rồi mask)

Package check và install do OpenCode Agent thực hiện qua bash trước khi gọi script này.
"""

import re
import json
import argparse
import sys
from pathlib import Path

# ======= VERSION =======
__version__ = "1.9.0"

# ======= PATTERNS =======
_HO_VIET = (
    r'(?:Nguyễn|NGUYỄN|Trần|TRẦN|Lê|LÊ|Phạm|PHẠM|Hoàng|HOÀNG|Huỳnh|HUỲNH|'
    r'Phan|PHAN|Vũ|VŨ|Võ|VÕ|Đặng|ĐẶNG|Bùi|BÙI|Đỗ|ĐỖ|'
    r'Dương|DƯƠNG|Đinh|ĐINH|Lưu|LƯU|Đào|ĐÀO|Tạ|TẠ|Vương|VƯƠNG|Lý|LÝ|Triệu|TRIỆU|'
    r'Quách|QUÁCH|Nông|NÔNG|Lục|LỤC|Tống|TỐNG|Lâm|LÂM|'
    r'Trịnh|TRỊNH|Đoàn|ĐOÀN|Thạch|THẠCH|Khúc|KHÚC|Khổng|KHỔNG|Khương|KHƯƠNG|'
    r'Châu|CHÂU|Chiêm|CHIÊM|Dư|DƯ|Diệp|DIỆP|Giáp|GIÁP|'
    r'Hứa|HỨA|Hàn|HÀN|Hàm|HÀM|'
    r'Kiều|KIỀU|Khuất|KHUẤT|Lạc|LẠC|Mạc|MẠC|Mạch|MẠCH|Mã|MÃ|'
    r'Nhan|NHAN|Nghiêm|NGHIÊM|Ninh|NINH|Ông|ÔNG|'
    r'Sầm|SẦM|Ứng|ỨNG|Úc|ÚC|Xa|XA|Xào|XÀO|Yên|YÊN|'
    r'Hà|HÀ|Mai|MAI|'
    r'(?<!Chí )(?<!CHÍ )Ngô|(?<!Chí )(?<!CHÍ )NGÔ|'
    r'(?<!Chí )(?<!CHÍ )Hồ(?!\s+Chí)(?!\s+CHÍ)|'
    r'(?<!Chí )(?<!CHÍ )HỒ(?!\s+Chí)(?!\s+CHÍ))'
)
_HO_NODIAC = (
    r'(?:NGUYEN|TRAN|LE|PHAM|HOANG|HUYNH|PHAN|VU|VO|DANG|BUI|DO|HO|NGO|DUONG|'
    r'DINH|LUU|DAO|TA|VUONG|LY|QUACH|NONG|LUC|TONG|LAM|TRIEU|'
    r'TRINH|DOAN|THACH|KHUC|KHONG|KHUONG|'
    r'CHAU|CHIEM|DU|DIEP|GIAP|HUA|HAN|HAM|'
    r'KIEU|KHUAT|LAC|MAC|MACH|MA|NHAN|NGHIEM|NINH|ONG|'
    r'SAM|UNG|UC|YEN|HA|MAI|'
    r'Nguyen|Tran|Le|Pham|Hoang|Huynh|Phan|Vu|Vo|Dang|Bui|Do|Ho|Ngo|Duong|'
    r'Dinh|Luu|Dao|Ta|Vuong|Ly|Quach|Nong|Luc|Tong|Lam|Trieu|'
    r'Trinh|Doan|Thach|Khuc|Khong|Khuong|'
    r'Chau|Chiem|Du|Diep|Giap|Hua|Han|Ham|'
    r'Kieu|Khuat|Lac|Mac|Mach|Ma|Nhan|Nghiem|Ninh|Ong|'
    r'Sam|Ung|Uc|Yen|Ha|Mai)'
    # Đã loại: Cao/CAO, Ly/LY, To/TO, Dong/DONG, Tien/TIEN, Cu/CU, Cai/CAI, Cam/CAM
)
_WORD_UPPER  = r'[A-ZÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚÝĂĐƠƯẢẤẦẨẪẬẮẰẲẴẶẺẼẾỀỂỄỆỈỊỎỐỒỔỖỘỚỜỞỠỢỤỦỨỪỬỮỰỲỴỶỸ][^\s,.:;!?()\[\]{}<>0-9]{1,15}'
_WORD_NODIAC = r'[A-Z][A-Za-z]{1,15}'

_DEPT_PREFIXES = (
    r'(?:Phòng|PHÒNG|Bộ phận|BỘ PHẬN|Khối|KHỐI|Trung tâm|TRUNG TÂM|'
    r'Văn phòng|VĂN PHÒNG|Chi nhánh|CHI NHÁNH|Tổ|TỔ|Nhóm|NHÓM|'
    r'Team|TEAM|Dept\.?|DEPT\.?|Department|DEPARTMENT|Division|DIVISION|'
    r'Center|CENTER|Office|OFFICE|'
    r'Ban(?=\s+[A-ZÀÁÂÃĐÈÉÊÌÍÒÓÔÕÙÚÝĂƠƯ])|'
    r'BAN(?=\s+[A-ZÀÁÂÃĐÈÉÊÌÍÒÓÔÕÙÚÝĂƠƯ]))'
)
_DEPT_EXCLUDE = r'(?!(?:họp|đầu|này|đó|kia|nào|nay|trống|vệ sinh|tắm|ăn|ngủ|chờ|đợi)\b)'
# Chỉ match từ bắt đầu chữ hoa, KHÔNG phải viết tắt quý (Q1-Q4), năm (2020-2029)
_DEPT_WORD    = r'(?!Q[1-4]\b)(?!20\d\d\b)[A-ZÀÁÂÃĐÈÉÊÌÍÒÓÔÕÙÚÝĂƠƯ][A-Za-zÀ-ỹ\.]*'

PATTERNS = {
    "phone":   (
        r'(?<!\d)'
        r'(\+84[\s\-\.]?|0)'
        r'(3[2-9]|5[6-9]|7[06-9]|8[0-9]|9[0-9])'
        r'([\s\-\.]?\d){7}'
        r'(?!\d)'
    ),
    "email":   r'[\w.\-]+@[\w.\-]+\.[a-zA-Z]{2,}',
    "cccd":    r'(?<!\d)\d{9}(?!\d)|(?<!\d)\d{12}(?!\d)',
    "account": r'(?<!\d)\d{10,16}(?!\d)',
    "salary":  r'(?<!\d)\d{1,3}(?:[.,]\d{3})*\s*(?:triệu|tr|VND|đồng|vnđ)(?!\w)',
    # Tên có dấu: dùng IGNORECASE → match cả hoa/thường/mixed
    "name":    rf'(?<!\w){_HO_VIET}(?:\s+{_WORD_UPPER}){{1,3}}(?!\w)',
    # Tên không dấu: KHÔNG dùng IGNORECASE → chỉ match Title Case hoặc UPPER CASE
    "name_nodiac": rf'(?<!\w){_HO_NODIAC}(?:\s+{_WORD_NODIAC}){{1,3}}(?!\w)',
    # Phòng ban: whitelist prefix + 1-4 từ theo sau
    "dept":    rf'(?<!\w){_DEPT_PREFIXES}\s+{_DEPT_EXCLUDE}(?:{_DEPT_WORD}\s*){{1,4}}',
}


# ======= MAPPING MANAGER =======
class MappingManager:
    def __init__(self):
        self.mapping = {}  # token → original
        self.reverse = {}  # original → token
        self.counters = {}  # type → count

    def get_or_create(self, data_type, value):
        key = f"{data_type}::{value}"
        if key not in self.reverse:
            self.counters[data_type] = self.counters.get(data_type, 0) + 1
            n = self.counters[data_type]
            label = {
                "phone":       "SĐT",
                "email":       "EMAIL",
                "cccd":        "ID",
                "account":     "TKNH",
                "name":        "TÊN",
                "name_nodiac": "TÊN",
                "salary":      "LƯƠNG",
                "dept":        "PHÒNG",
                "dept_custom": "PHÒNG",
            }.get(data_type, data_type.upper())
            token = f"[{label}_{n:03d}]"
            self.mapping[token] = value
            self.reverse[key] = token
        return self.reverse[key]

    def summary(self):
        return {t: c for t, c in self.counters.items()}

    def save(self, path):
        with open(path, "w", encoding="utf-8") as f:
            json.dump(
                {"mapping": self.mapping, "summary": self.summary()},
                f,
                ensure_ascii=False,
                indent=2,
            )



# Các loại chứa số — bị loại khi --keep-numbers
NUMBER_TYPES = {"account", "salary"}


# ======= MASK TEXT =======
def mask_text(text, mgr, types="all", keep_numbers=False):
    if not text or not isinstance(text, str):
        return text

    active = list(PATTERNS.keys()) if types == "all" else types.split(",")

    # --keep-numbers: bỏ qua pattern liên quan đến số
    if keep_numbers:
        active = [t for t in active if t not in NUMBER_TYPES]

    for data_type in active:
        if data_type not in PATTERNS:
            continue
        pattern = PATTERNS[data_type]
        def replacer(m, dt=data_type):
            # name_nodiac và dept_custom map về loại gốc để dùng chung counter
            key = {
                "name_nodiac": "name",
                "dept_custom":  "dept",
            }.get(dt, dt)
            return mgr.get_or_create(key, m.group(0).strip())
        # Các loại KHÔNG dùng IGNORECASE — tên người và phòng ban phải viết hoa
        no_ignore = {"name", "name_nodiac", "dept", "dept_custom"}
        flags = 0 if data_type in no_ignore else re.IGNORECASE
        text = re.sub(pattern, replacer, text, flags=flags)

    return text


# ======= FILE HANDLERS =======
def process_txt(input_path, output_path, mgr, types, keep_numbers=False):
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()
    masked = mask_text(content, mgr, types, keep_numbers)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(masked)


def process_csv(input_path, output_path, mgr, types, keep_numbers=False):
    import csv

    rows = []
    with open(input_path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append([mask_text(cell, mgr, types, keep_numbers) for cell in row])
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(rows)


def process_docx(input_path, output_path, mgr, types, keep_numbers=False):
    from docx import Document

    doc = Document(input_path)
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = mask_text(run.text, mgr, types, keep_numbers)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = mask_text(run.text, mgr, types, keep_numbers)
    doc.save(output_path)


def process_xlsx(input_path, output_path, mgr, types, keep_numbers=False):
    import openpyxl

    wb = openpyxl.load_workbook(input_path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value = mask_text(cell.value, mgr, types, keep_numbers)
                # Cell kiểu số (int/float) → giữ nguyên hoàn toàn, không đụng vào
    wb.save(output_path)


def process_pdf(input_path, output_path, mgr, types, keep_numbers=False):
    from pdf2docx import Converter

    output_path = str(Path(output_path).with_suffix(".docx"))
    print(f"  [PDF] Đang convert {input_path} → {output_path} ...")
    cv = Converter(input_path)
    cv.convert(output_path, start=0, end=None)
    cv.close()
    print(f"  [PDF] Convert xong, đang mask...")
    process_docx(output_path, output_path, mgr, types, keep_numbers)


# ======= REVERSE HANDLERS =======
def reverse_text(text, mapping):
    """Thay thế tất cả token [XXX_NNN] về giá trị gốc từ mapping."""
    if not text or not isinstance(text, str):
        return text
    for token, original in mapping.items():
        text = text.replace(token, original)
    return text


def reverse_txt(input_path, output_path, mapping):
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(reverse_text(content, mapping))


def reverse_csv(input_path, output_path, mapping):
    import csv

    rows = []
    with open(input_path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append([reverse_text(cell, mapping) for cell in row])
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(rows)


def reverse_docx(input_path, output_path, mapping):
    from docx import Document

    doc = Document(input_path)
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = reverse_text(run.text, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = reverse_text(run.text, mapping)
    doc.save(output_path)


def reverse_xlsx(input_path, output_path, mapping):
    import openpyxl

    wb = openpyxl.load_workbook(input_path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value = reverse_text(cell.value, mapping)
    wb.save(output_path)


REVERSE_HANDLERS = {
    ".txt": reverse_txt,
    ".csv": reverse_csv,
    ".docx": reverse_docx,
    ".xlsx": reverse_xlsx,
}


HANDLERS = {
    ".txt": process_txt,
    ".csv": process_csv,
    ".docx": process_docx,
    ".xlsx": process_xlsx,
    ".pdf": process_pdf,
}


# ======= MAIN =======
def main():
    parser = argparse.ArgumentParser(description="Data Anonymity — local anonymization")
    parser.add_argument("--version", action="store_true")
    parser.add_argument("--input", "-i", help="File đầu vào")
    parser.add_argument("--output", "-o", help="File đầu ra (đã mask)")
    parser.add_argument(
        "--mapping", "-m", help="File mapping JSON (bắt buộc khi dùng --reverse)"
    )
    parser.add_argument(
        "--types",
        "-t",
        default="all",
        help="Loại dữ liệu: all | phone,email,cccd,account,name,salary,dept",
    )
    parser.add_argument(
        "--keep-numbers",
        action="store_true",
        help="Giữ nguyên mọi số, chỉ mask text (phù hợp dữ liệu tài chính)",
    )
    parser.add_argument(
        "--departments", "-d",
        help="Danh sách tên phòng ban cụ thể, cách nhau bằng dấu phẩy. VD: 'R&D,Marketing,C-Suite'",
    )
    parser.add_argument(
        "--reverse",
        action="store_true",
        help="Reverse file đã mask về dữ liệu gốc, cần --mapping",
    )
    parser.add_argument("--summary", "-s", help="In summary từ file mapping JSON")
    args = parser.parse_args()

    # Inject custom department pattern nếu có --departments
    if args.departments:
        custom_depts = [d.strip() for d in args.departments.split(",") if d.strip()]
        if custom_depts:
            escaped = [re.escape(d) for d in custom_depts]
            PATTERNS["dept_custom"] = r'(?<!\w)(?:' + '|'.join(escaped) + r')(?!\w)'
            print(f"  [dept] Thêm {len(custom_depts)} phòng ban tùy chỉnh: {', '.join(custom_depts)}")
    args = parser.parse_args()

    if args.version:
        print(f"data-anonymity v{__version__}")
        return

    if args.summary:
        with open(args.summary, "r", encoding="utf-8") as f:
            data = json.load(f)
        print("=== KẾT QUẢ ANONYMIZE ===")
        for k, v in data.get("summary", {}).items():
            print(f"  {k:10s}: {v} giá trị")
        print("=========================")
        return

    if not args.input:
        print("Lỗi: cần --input <file>", file=sys.stderr)
        sys.exit(1)

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Lỗi: file không tồn tại: {input_path}", file=sys.stderr)
        sys.exit(1)

    ext = input_path.suffix.lower()

    # ── REVERSE MODE ──────────────────────────────────────────────────────────
    if args.reverse:
        if not args.mapping:
            print("Lỗi: --reverse cần --mapping <file_mapping.json>", file=sys.stderr)
            sys.exit(1)
        if ext not in REVERSE_HANDLERS:
            print(
                f"Lỗi: định dạng chưa hỗ trợ reverse: {ext}. Hỗ trợ: {', '.join(REVERSE_HANDLERS)}",
                file=sys.stderr,
            )
            sys.exit(1)

        mapping_path = Path(args.mapping)
        if not mapping_path.exists():
            print(f"Lỗi: file mapping không tồn tại: {mapping_path}", file=sys.stderr)
            sys.exit(1)

        with open(mapping_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        mapping = data.get("mapping", {})

        output_path = (
            Path(args.output)
            if args.output
            else input_path.with_stem(input_path.stem + "_restored")
        )

        print(f"Đang reverse: {input_path}")
        print(f"  Mapping: {mapping_path} ({len(mapping)} token)")
        REVERSE_HANDLERS[ext](str(input_path), str(output_path), mapping)
        print(f"✓ Restored: {output_path}")
        return

    # ── ANONYMIZE MODE ────────────────────────────────────────────────────────
    if ext not in HANDLERS:
        print(
            f"Lỗi: định dạng chưa hỗ trợ: {ext}. Hỗ trợ: {', '.join(HANDLERS)}",
            file=sys.stderr,
        )
        sys.exit(1)

    output_path = (
        Path(args.output)
        if args.output
        else input_path.with_stem(input_path.stem + "_masked")
    )
    if ext == ".pdf":
        output_path = output_path.with_suffix(".docx")
    mapping_path = (
        Path(args.mapping)
        if args.mapping
        else input_path.with_stem(input_path.stem + "_mapping").with_suffix(".json")
    )

    mgr = MappingManager()

    keep_numbers = args.keep_numbers
    if keep_numbers:
        print(f"  [chế độ Finance] Giữ nguyên số liệu, chỉ mask định danh")

    print(f"Đang xử lý: {input_path}")
    HANDLERS[ext](str(input_path), str(output_path), mgr, args.types, keep_numbers)
    mgr.save(str(mapping_path))

    print(f"✓ Output:  {output_path}")
    print(f"✓ Mapping: {mapping_path}  ← giữ file này để reverse nếu cần")
    print(f"✓ Tổng kết:")
    for k, v in mgr.summary().items():
        print(f"  {k:10s}: {v} giá trị đã mask")


if __name__ == "__main__":
    main()
